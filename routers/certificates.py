from fastapi import APIRouter, Request, Depends, Form, UploadFile, File
from fastapi.responses import HTMLResponse, RedirectResponse, StreamingResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from database import get_db
from models.models import CertificateTemplate, EventParticipant, Event, Student, CustomTemplate, Admin
from routers.auth import verify_token
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas as pdf_canvas
from reportlab.lib import colors
from fastapi.responses import Response
import io
import os
import uuid
import base64
import tempfile
import storage_fb

router = APIRouter()
from templating import templates

UPLOAD_DIR = "static/templates_upload"
os.makedirs(UPLOAD_DIR, exist_ok=True)


def _storage_path(tmpl) -> str:
    return f"templates/{tmpl.filename}"


def _media_type_for(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    return {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }.get(ext, "application/octet-stream")


def _template_bytes(tmpl) -> bytes:
    """Raw bytes from Firebase Storage (new), legacy base64, or local disk."""
    if getattr(tmpl, "file_data", None):
        return base64.b64decode(tmpl.file_data)  # legacy rows stored in Firestore
    try:
        return storage_fb.download_bytes(_storage_path(tmpl))
    except Exception:
        with open(os.path.join(UPLOAD_DIR, tmpl.filename), "rb") as f:
            return f.read()


def _template_to_tempfile(tmpl) -> str:
    """Write the template to a temp file and return its path (caller deletes it)."""
    ext = os.path.splitext(tmpl.filename)[1].lower()
    fd, path = tempfile.mkstemp(suffix=ext)
    with os.fdopen(fd, "wb") as f:
        f.write(_template_bytes(tmpl))
    return path

@router.get("/certificates", response_class=HTMLResponse)
async def certificates_page(request: Request, db: Session = Depends(get_db)):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    cert_templates = db.query(CertificateTemplate).all()
    custom_templates = db.query(CustomTemplate).all()
    events = db.query(Event).filter(Event.is_completed == True).all()
    return templates.TemplateResponse(request, "admin/certificates.html", {
        "cert_templates": cert_templates,
        "custom_templates": custom_templates,
        "events": events
    })

@router.get("/certificates/template/edit/{template_id}", response_class=HTMLResponse)
async def edit_template_page(template_id: str, request: Request, db: Session = Depends(get_db)):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    tmpl = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not tmpl:
        return RedirectResponse(url="/certificates", status_code=303)
    return templates.TemplateResponse(request, "admin/certificate_editor.html", {
        "tmpl": tmpl,
        "active": "certificates"
    })

@router.post("/certificates/template/edit/{template_id}")
async def save_template_settings(
    template_id: str,
    request: Request,
    db: Session = Depends(get_db)
):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    form = await request.form()
    tmpl = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if tmpl:
        tmpl.name_x = int(form.get("name_x", 50))
        tmpl.name_y = int(form.get("name_y", 45))
        tmpl.name_font_size = int(form.get("name_font_size", 36))
        tmpl.name_color = form.get("name_color", "#1a2535")
        tmpl.event_x = int(form.get("event_x", 50))
        tmpl.event_y = int(form.get("event_y", 58))
        tmpl.event_font_size = int(form.get("event_font_size", 20))
        tmpl.event_color = form.get("event_color", "#374151")
        tmpl.position_x = int(form.get("position_x", 50))
        tmpl.position_y = int(form.get("position_y", 65))
        tmpl.position_font_size = int(form.get("position_font_size", 24))
        tmpl.position_color = form.get("position_color", "#b45309")
        tmpl.house_x = int(form.get("house_x", 50))
        tmpl.house_y = int(form.get("house_y", 72))
        tmpl.house_font_size = int(form.get("house_font_size", 16))
        tmpl.house_color = form.get("house_color", "#374151")
        tmpl.date_x = int(form.get("date_x", 50))
        tmpl.date_y = int(form.get("date_y", 78))
        tmpl.date_font_size = int(form.get("date_font_size", 14))
        tmpl.date_color = form.get("date_color", "#6b7280")
        db.commit()
    return RedirectResponse(url="/certificates?msg=settings_saved", status_code=303)

@router.post("/certificates/template/save")
async def save_builtin_template(
    request: Request,
    name: str = Form(...),
    title_text: str = Form(...),
    body_text: str = Form(...),
    font_family: str = Form(...),
    is_default: bool = Form(False),
    db: Session = Depends(get_db)
):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    if is_default:
        for t in db.query(CertificateTemplate).all():
            t.is_default = False
    template = CertificateTemplate(
        name=name,
        title_text=title_text,
        body_text=body_text,
        font_family=font_family,
        is_default=is_default
    )
    db.add(template)
    db.commit()
    return RedirectResponse(url="/certificates?msg=saved", status_code=303)

@router.post("/certificates/template/upload")
async def upload_custom_template(
    request: Request,
    name: str = Form(...),
    template_file: UploadFile = File(...),
    is_default: bool = Form(False),
    db: Session = Depends(get_db)
):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    ext = os.path.splitext(template_file.filename)[1].lower()
    if ext not in [".docx", ".jpeg", ".jpg", ".png"]:
        return RedirectResponse(url="/certificates?msg=invalid_file", status_code=303)
    raw = await template_file.read()
    file_type = "docx" if ext == ".docx" else "image"
    filename = f"{uuid.uuid4()}{ext}"
    # Store the file in Firebase Storage (durable, no size cap).
    try:
        storage_fb.upload_bytes(f"templates/{filename}", raw, _media_type_for(filename))
    except Exception as e:
        print("[storage] template upload failed:", e)
        return RedirectResponse(url="/certificates?msg=upload_failed", status_code=303)
    if is_default:
        for t in db.query(CustomTemplate).all():
            t.is_default = False
    custom = CustomTemplate(
        name=name,
        filename=filename,
        file_type=file_type,
        is_default=is_default,
    )
    db.add(custom)
    db.commit()
    return RedirectResponse(url=f"/certificates/template/edit/{custom.id}", status_code=303)


@router.get("/certificates/template/file/{template_id}")
async def template_file(template_id: str, request: Request, db: Session = Depends(get_db)):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    tmpl = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if not tmpl:
        return Response(status_code=404)
    try:
        return Response(content=_template_bytes(tmpl), media_type=_media_type_for(tmpl.filename))
    except Exception:
        return Response(status_code=404)

@router.post("/certificates/template/delete/{template_id}")
async def delete_custom_template(
    template_id: str,
    request: Request,
    db: Session = Depends(get_db)
):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    tmpl = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if tmpl:
        if not getattr(tmpl, "file_data", None):
            storage_fb.delete_blob(f"templates/{tmpl.filename}")
        filepath = os.path.join(UPLOAD_DIR, tmpl.filename)
        if os.path.exists(filepath):
            os.remove(filepath)
        db.delete(tmpl)
        db.commit()
    return RedirectResponse(url="/certificates?msg=deleted", status_code=303)

@router.post("/certificates/template/set_default/{template_id}")
async def set_default_custom_template(
    template_id: str,
    request: Request,
    db: Session = Depends(get_db)
):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    for t in db.query(CustomTemplate).all():
        t.is_default = False
    tmpl = db.query(CustomTemplate).filter(CustomTemplate.id == template_id).first()
    if tmpl:
        tmpl.is_default = True
    db.commit()
    return RedirectResponse(url="/certificates?msg=default_set", status_code=303)

POSITION_LABELS = {1: "1st", 2: "2nd", 3: "3rd"}


def _pos_label(pos):
    return POSITION_LABELS.get(pos, f"{pos}th")


def _winners_for(db, event_id):
    return db.query(EventParticipant).filter(
        EventParticipant.event_id == event_id,
        EventParticipant.position != None  # noqa: E711
    ).order_by(EventParticipant.position).all()


def _dispatch(pairs, db, download_name):
    """Render a single file for a list of (winner, event) pairs, using whichever
    default template is configured (custom image/docx, else the built-in PDF)."""
    custom_template = db.query(CustomTemplate).filter(CustomTemplate.is_default == True).first()
    if custom_template:
        tmp_path = None
        try:
            tmp_path = _template_to_tempfile(custom_template)
        except Exception:
            tmp_path = None
        if tmp_path:
            try:
                if custom_template.file_type == "docx":
                    return generate_from_docx(tmp_path, pairs, download_name)
                return generate_from_image(tmp_path, pairs, custom_template, download_name)
            finally:
                try:
                    os.remove(tmp_path)
                except Exception:
                    pass
    return generate_default_pdf(pairs, db, download_name)


@router.get("/certificates/generate/{event_id}")
async def generate_certificates(
    event_id: str,
    request: Request,
    db: Session = Depends(get_db)
):
    if not verify_token(request):
        return RedirectResponse(url="/login", status_code=303)
    event = db.query(Event).filter(Event.id == event_id).first()
    if not event:
        return RedirectResponse(url="/certificates", status_code=303)
    winners = _winners_for(db, event_id)
    if not winners:
        return RedirectResponse(url="/certificates?msg=no_winners", status_code=303)
    pairs = [(w, event) for w in winners]
    return _dispatch(pairs, db, f"certificates_{event.name.replace(' ', '_')}")

def _is_super(request, db):
    email = verify_token(request)
    if not email:
        return None
    admin = db.query(Admin).filter(Admin.email == email).first()
    return admin if (admin and admin.role == "super_admin") else None


@router.get("/certificates/bulk", response_class=HTMLResponse)
async def bulk_page(request: Request, db: Session = Depends(get_db)):
    if not _is_super(request, db):
        return RedirectResponse(url="/dashboard", status_code=303)
    events = db.query(Event).filter(Event.is_completed == True).order_by(Event.event_date.desc()).all()
    rows = []
    for ev in events:
        cnt = db.query(EventParticipant).filter(
            EventParticipant.event_id == ev.id,
            EventParticipant.position != None  # noqa: E711
        ).count()
        rows.append({"event": ev, "winners": cnt})
    return templates.TemplateResponse(request, "admin/bulk_certificates.html", {
        "active": "bulk_certs",
        "rows": rows,
    })


@router.post("/certificates/bulk")
async def bulk_generate(request: Request, db: Session = Depends(get_db)):
    if not _is_super(request, db):
        return RedirectResponse(url="/dashboard", status_code=303)
    form = await request.form()
    ids = form.getlist("event_ids")
    pairs = []
    for eid in ids:
        ev = db.query(Event).filter(Event.id == eid).first()
        if not ev:
            continue
        for w in _winners_for(db, eid):
            pairs.append((w, ev))
    if not pairs:
        return RedirectResponse(url="/certificates/bulk?msg=no_winners", status_code=303)
    return _dispatch(pairs, db, "certificates_bulk")


def generate_from_docx(filepath, pairs, download_name="certificates"):
    from docx import Document
    import copy
    buffer = io.BytesIO()
    merged_doc = Document()
    for i, (winner, event) in enumerate(pairs):
        student = winner.student
        pos_label = _pos_label(winner.position)
        replacements = {
            "{{name}}": student.name,
            "{{event}}": event.name,
            "{{position}}": pos_label,
            "{{house}}": student.house.name,
            "{{class}}": student.class_name,
            "{{date}}": str(event.event_date),
            "{{school}}": "National Public School",
        }
        temp_doc = Document(filepath)
        for para in temp_doc.paragraphs:
            for run in para.runs:
                for key, val in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)
        for table in temp_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for key, val in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, val)
        if i == 0:
            for element in temp_doc.element.body:
                merged_doc.element.body.append(copy.deepcopy(element))
        else:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            page_break = OxmlElement('w:p')
            pb_run = OxmlElement('w:r')
            pb = OxmlElement('w:br')
            pb.set(qn('w:type'), 'page')
            pb_run.append(pb)
            page_break.append(pb_run)
            merged_doc.element.body.append(page_break)
            for element in temp_doc.element.body:
                merged_doc.element.body.append(copy.deepcopy(element))
    merged_doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={download_name}.docx"}
    )

def generate_from_image(filepath, pairs, tmpl, download_name="certificates"):
    from PIL import Image as PILImage
    buffer = io.BytesIO()

    # Detect image size and use as page size
    with PILImage.open(filepath) as img:
        img_width, img_height = img.size

    # Convert pixels to points (72 points per inch, assume 96 DPI)
    pt_width = img_width * 72 / 96
    pt_height = img_height * 72 / 96

    c = pdf_canvas.Canvas(buffer, pagesize=(pt_width, pt_height))

    for winner, event in pairs:
        student = winner.student
        pos_label = _pos_label(winner.position)

        # Draw background image filling entire page
        c.drawImage(filepath, 0, 0, width=pt_width, height=pt_height)

        def hex_to_rgb(hex_color):
            hex_color = hex_color.lstrip('#')
            return tuple(int(hex_color[i:i+2], 16)/255 for i in (0, 2, 4))

        def draw_text(text, x_pct, y_pct, font_size, color_hex):
            r, g, b = hex_to_rgb(color_hex)
            c.setFillColorRGB(r, g, b)
            c.setFont("Helvetica-Bold", font_size)
            # x_pct and y_pct are percentages of page dimensions
            x = pt_width * x_pct / 100
            # PDF y starts from bottom, so invert
            y = pt_height * (100 - y_pct) / 100
            c.drawCentredString(x, y, text)

        # Draw each field using saved positions
        draw_text(student.name, tmpl.name_x, tmpl.name_y,
                  tmpl.name_font_size, tmpl.name_color)
        draw_text(f"{pos_label} Place — {event.name}", tmpl.event_x, tmpl.event_y,
                  tmpl.event_font_size, tmpl.event_color)
        draw_text(pos_label, tmpl.position_x, tmpl.position_y,
                  tmpl.position_font_size, tmpl.position_color)
        draw_text(f"House: {student.house.name} | Class: {student.class_name}",
                  tmpl.house_x, tmpl.house_y,
                  tmpl.house_font_size, tmpl.house_color)
        draw_text(str(event.event_date), tmpl.date_x, tmpl.date_y,
                  tmpl.date_font_size, tmpl.date_color)

        c.showPage()

    c.save()
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={download_name}.pdf"}
    )

def generate_default_pdf(pairs, db, download_name="certificates"):
    template = db.query(CertificateTemplate).filter(
        CertificateTemplate.is_default == True
    ).first()
    if not template:
        template = CertificateTemplate(
            title_text="Certificate of Achievement",
            body_text="This is to certify that {name} has achieved {position} place in {event} at National Public School.",
            font_family="Helvetica"
        )
    buffer = io.BytesIO()
    page_width, page_height = landscape(A4)
    c = pdf_canvas.Canvas(buffer, pagesize=landscape(A4))
    for winner, event in pairs:
        student = winner.student
        pos_label = _pos_label(winner.position)
        body = template.body_text.replace("{name}", student.name)
        body = body.replace("{event}", event.name)
        body = body.replace("{position}", pos_label)
        c.setFillColorRGB(0.98, 0.97, 0.93)
        c.rect(0, 0, page_width, page_height, fill=1, stroke=0)
        c.setStrokeColorRGB(0.72, 0.58, 0.22)
        c.setLineWidth(3)
        c.rect(30, 30, page_width-60, page_height-60, fill=0, stroke=1)
        c.setLineWidth(1)
        c.rect(40, 40, page_width-80, page_height-80, fill=0, stroke=1)
        c.setFillColorRGB(0.15, 0.15, 0.35)
        c.setFont("Helvetica-Bold", 13)
        c.drawCentredString(page_width/2, page_height-80, "NATIONAL PUBLIC SCHOOL")
        c.setStrokeColorRGB(0.72, 0.58, 0.22)
        c.setLineWidth(1.5)
        c.line(page_width/2-180, page_height-95, page_width/2+180, page_height-95)
        c.setFillColorRGB(0.55, 0.40, 0.05)
        c.setFont("Helvetica-Bold", 36)
        c.drawCentredString(page_width/2, page_height-150, template.title_text.upper())
        c.setFillColorRGB(0.3, 0.3, 0.3)
        c.setFont("Helvetica", 13)
        c.drawCentredString(page_width/2, page_height-180, "This certificate is proudly presented to")
        c.setFillColorRGB(0.1, 0.15, 0.35)
        c.setFont("Helvetica-Bold", 30)
        c.drawCentredString(page_width/2, page_height-230, student.name)
        name_width = c.stringWidth(student.name, "Helvetica-Bold", 30)
        c.setStrokeColorRGB(0.72, 0.58, 0.22)
        c.setLineWidth(1)
        c.line(page_width/2-name_width/2, page_height-238, page_width/2+name_width/2, page_height-238)
        c.setFillColorRGB(0.25, 0.25, 0.25)
        c.setFont("Helvetica", 14)
        words = body.split()
        lines = []
        current_line = ""
        for word in words:
            test = current_line + " " + word if current_line else word
            if c.stringWidth(test, "Helvetica", 14) < page_width - 200:
                current_line = test
            else:
                lines.append(current_line)
                current_line = word
        if current_line:
            lines.append(current_line)
        y = page_height - 275
        for line in lines:
            c.drawCentredString(page_width/2, y, line)
            y -= 22
        c.setFillColorRGB(0.45, 0.45, 0.45)
        c.setFont("Helvetica", 11)
        c.drawCentredString(page_width/2, page_height-340,
            f"House: {student.house.name}   |   Class: {student.class_name}   |   Date: {event.event_date}")
        c.setFillColorRGB(0.72, 0.58, 0.22)
        c.setFont("Helvetica-Bold", 16)
        c.drawCentredString(page_width/2, page_height-370, f"{pos_label} Place — {event.name}")
        c.setStrokeColorRGB(0.72, 0.58, 0.22)
        c.setLineWidth(1.5)
        c.line(page_width/2-180, page_height-390, page_width/2+180, page_height-390)
        c.setStrokeColorRGB(0.5, 0.5, 0.5)
        c.setLineWidth(0.8)
        c.line(120, 80, 280, 80)
        c.line(page_width-280, 80, page_width-120, 80)
        c.setFillColorRGB(0.4, 0.4, 0.4)
        c.setFont("Helvetica", 10)
        c.drawCentredString(200, 68, "Principal")
        c.drawCentredString(page_width-200, 68, "Class Teacher")
        c.showPage()
    c.save()
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={download_name}.pdf"}
    )